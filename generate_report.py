#!/usr/bin/env python3
"""
Word report generator:
  python generate_report.py --format word --output E_Learning_Project_Report.docx
"""
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_code(doc, code: str):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(12)


def build_document():
    doc = Document()
    sect = doc.sections[0]
    footer = sect.footer.paragraphs[0] if sect.footer.paragraphs else sect.footer.add_paragraph()
    footer.text = "E-Learning Platform — Project Report"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph()
    r = title.add_run("E-Learning Platform (Django REST + React)")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Technical & learning documentation").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_heading(doc, "Table of Contents", level=1)
    add_para(
        doc,
        "Word mein TOC auto page numbers ke liye: References → Table of Contents. "
        "Is generated file mein humne manual outline rakha hai taake script simple rahe.",
    )
    for i in range(1, 11):
        doc.add_paragraph(f"{i}. Chapter {i} (see headings below)", style="List Number")

    chapters = [
        (
            "Chapter 1: Project Overview",
            [
                "Project name: E-Learning Platform.",
                "Description: Online courses, video lessons, quizzes, Stripe payments, reviews, notifications, realtime chat (Channels).",
                "Audience: Students, instructors, admins.",
                "Business value: Marketplace-style discovery + structured learning + monetization hooks.",
                "Competitive angle: Similar surface area to Udemy (course commerce) and Coursera-like structure (curriculum/progress), implemented as an API-first stack.",
            ],
        ),
        (
            "Chapter 2: Technical Architecture",
            [
                "Browser (React/Vite) talks to Django via `/api/v1` (dev proxy).",
                "PostgreSQL (prod) / SQLite (local demo) stores relational data.",
                "Redis: cache + channel layers + Celery broker (recommended).",
                "Async: Celery tasks for PDF certificate generation; Channels ASGI for WebSockets.",
                "API docs: `/api/docs/` using drf-spectacular.",
            ],
        ),
        (
            "Chapter 3: Database Design",
            [
                "ER (text): User 1—* Enrollment *—1 Course; Course 1—* Lesson; Lesson M2M prerequisites; Course 1—* Review; Order(User,Course); Quiz→Question→Choice.",
                "Indexes: course filters (status/pricing), ratings, enrollments_count; lesson ordering per course.",
            ],
        ),
        (
            "Chapter 4: Module-wise Implementation",
            [
                "Authentication: simplejwt + register + signed email verify + password reset + optional Google token verification.",
                "Courses: CRUD, filters/search/ordering, enroll, progress, bookmarks, notes, announcements, certificate PDF task.",
                "Payments: Stripe Checkout + webhook; coupons; demo paid flow without keys.",
                "Quizzes: MCQ auto scoring; coding submissions stub autograde; peer review model.",
                "Reviews: aggregates via signals; instructor reply; helpful votes.",
                "Realtime chat: `ws/chat/<room>/` consumer + persisted messages REST read.",
                "Search: DRF SearchFilter (upgrade path: Elasticsearch).",
                "Admin analytics: `/api/v1/admin/platform-stats/`.",
            ],
        ),
        (
            "Chapter 5: Frontend Implementation",
            [
                "React 18 + Vite + Tailwind; TanStack Query for server cache; Zustand persist for tokens; protected routes; RHF+Zod on auth pages; Chart.js + React Player demos.",
            ],
        ),
        (
            "Chapter 6: Testing Strategy",
            [
                "pytest-django + `config.settings.test` uses LocMem cache so throttles work without Redis.",
                "Run: `python -m pytest backend/tests -q` (ensure `DJANGO_SETTINGS_MODULE=config.settings.test`).",
            ],
        ),
        (
            "Chapter 7: Deployment Guide",
            [
                "Prereqs: Python 3.12+, Node 20+, PostgreSQL 15+, Redis 7+.",
                "Backend: `pip install -r backend/requirements.txt`, env from `backend/.env.example`, `python manage.py migrate`, `python seed.py`, `runserver`.",
                "Frontend: `npm install && npm run dev`.",
                "Docker: use provided `docker-compose.yml` for web+db+redis+worker wiring.",
            ],
        ),
        (
            "Chapter 8: Learning Documentation (Hinglish notes)",
            [
                "ORM: Python classes = DB tables; relations queries ko readable banate hain.",
                "DRF: Serializers validation + ViewSets + routers = fast CRUD APIs.",
                "JWT: Stateless auth; refresh rotation blacklist se logout meaningful hota hai.",
                "Channels: WebSockets alag lifecycle; `group_send` broadcast pattern common hai.",
                "React Query: caching/refetch server state; UI fast feel hota hai.",
            ],
        ),
        (
            "Chapter 9: Troubleshooting",
            [
                "Redis errors: start Redis or use test settings for CI.",
                "Stripe webhook: signature secret mismatch = 400 (expected).",
                "CORS: ensure `CORS_ALLOWED_ORIGINS` includes Vite origin.",
            ],
        ),
        (
            "Chapter 10: Future Enhancements",
            [
                "AI recommendations, mobile app, live streaming, gamification.",
            ],
        ),
    ]

    for title_text, bullets in chapters:
        add_heading(doc, title_text, level=1)
        for b in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(b)

    add_heading(doc, "Appendix A: Sample endpoints", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Method"
    hdr[1].text = "URL"
    hdr[2].text = "Purpose"
    hdr[3].text = "Notes"
    rows = [
        ("POST", "/api/v1/auth/register/", "Signup + JWT", "email, password, role"),
        ("POST", "/api/v1/auth/login/", "JWT/cookies", "access+refresh"),
        ("GET", "/api/v1/courses/", "List", "filters + ordering"),
        ("POST", "/api/v1/courses/{slug}/enroll/", "Enroll", "paid needs PAID order"),
        ("POST", "/api/v1/checkout/create/", "Stripe checkout", "demo without keys"),
    ]
    for method, url, purpose, notes in rows:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = url
        row[2].text = purpose
        row[3].text = notes

    add_heading(doc, "Appendix B: Snippet — throttling rates", level=1)
    add_code(
        doc,
        "DEFAULT_THROTTLE_RATES = {\n"
        "  'anon': '120/hour',\n"
        "  'user': '2000/hour',\n"
        "  'login': '30/minute',\n"
        "  'burst': '60/minute',\n"
        "}",
    )

    return doc


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--format", default="word", choices=["word"])
    parser.add_argument("--output", default="E_Learning_Project_Report.docx")
    args = parser.parse_args()
    out = Path(args.output)
    doc = build_document()
    doc.save(out)
    print(f"Wrote {out.resolve()}")


if __name__ == "__main__":
    main()
